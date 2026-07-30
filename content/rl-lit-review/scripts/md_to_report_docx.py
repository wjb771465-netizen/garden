#!/usr/bin/env python3
"""Append Markdown (headings / prose / GFM tables) into report.docx.

Tables clone the visual style of the first table in the template
(header blue + alternating body rows, white borders).
"""

from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table

HEADER_FILL = "4F81BD"
BODY_FILLS = ("D0D8E8", "E9EDF4")
BORDER_COLOR = "FFFFFF"


_FRONTMATTER = re.compile(r"\A---\s*\n.*?\n---\s*\n?", re.DOTALL)


def parse_md(text: str) -> list[dict]:
    """Parse a small Markdown subset into blocks: heading / para / table."""
    text = _FRONTMATTER.sub("", text, count=1)
    lines = text.splitlines()
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if m := re.match(r"^(#{1,3})\s+(.+)$", line):
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue
        if _is_table_row(line) and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            header = _split_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                rows.append(_split_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not (
            _is_table_row(lines[i]) and i + 1 < len(lines) and _is_table_sep(lines[i + 1])
        ):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"type": "para", "text": " ".join(s.strip() for s in para_lines)})
    return blocks


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def _is_table_sep(line: str) -> bool:
    if not _is_table_row(line):
        return False
    cells = _split_row(line)
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells)


def _split_row(line: str) -> list[str]:
    inner = line.strip()[1:-1]
    return [c.strip() for c in inner.split("|")]


def _set_run_font(run, *, bold: bool = False, white: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(11)
    if white:
        run.font.color.rgb = None
        rPr = run._r.get_or_add_rPr()
        color = rPr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rPr.append(color)
        color.set(qn("w:val"), "FFFFFF")


def _clear_cell(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)
    tc.append(OxmlElement("w:p"))


def _set_cell_text(cell, text: str, *, header: bool = False) -> None:
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    run = p.add_run(text)
    _set_run_font(run, bold=header, white=header)


def _tc_borders(*, thick_top: bool = False, thick_bottom: bool = False) -> OxmlElement:
    borders = OxmlElement("w:tcBorders")
    specs = {
        "top": ("24" if thick_top else "8"),
        "left": "8",
        "bottom": ("24" if thick_bottom else "8"),
        "right": "8",
    }
    for edge, sz in specs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER_COLOR)
        borders.append(el)
    return borders


def _style_cell(cell, fill: str, *, header: bool = False, first_body: bool = False) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for tag in ("tcW", "tcBorders", "shd", "tcMar", "vAlign", "hideMark"):
        el = tcPr.find(qn(f"w:{tag}"))
        if el is not None:
            tcPr.remove(el)

    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), "0")
    tcW.set(qn("w:type"), "auto")
    tcPr.append(tcW)
    tcPr.append(_tc_borders(thick_top=first_body, thick_bottom=header))

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

    mar = OxmlElement("w:tcMar")
    for side, w in (("top", "54"), ("left", "108"), ("bottom", "54"), ("right", "108")):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), w)
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)
    tcPr.append(OxmlElement("w:hideMark"))


def _set_row_height(row, twips: int = 654) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    trHeight.set(qn("w:val"), str(twips))


def _clone_tbl_pr(src_tbl: Table, dst_tbl: Table) -> None:
    src_pr = src_tbl._tbl.tblPr
    if src_pr is None:
        return
    dst = dst_tbl._tbl
    old = dst.tblPr
    if old is not None:
        dst.remove(old)
    dst.insert(0, copy.deepcopy(src_pr))


def _set_grid(table: Table, col_widths: list[int]) -> None:
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    # tblGrid follows tblPr
    tblPr = tbl.tblPr
    idx = list(tbl).index(tblPr) + 1 if tblPr is not None else 0
    tbl.insert(idx, grid)


def _template_col_widths(template: Table, n_cols: int) -> list[int]:
    grid = template._tbl.tblGrid
    widths = [int(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol"))] if grid is not None else []
    if len(widths) == n_cols:
        return widths
    total = sum(widths) if widths else 10446
    if n_cols == 1:
        return [total]
    # keep first col narrow if template had a narrow index col and new table also starts with 编号
    base = [total // n_cols] * n_cols
    rem = total - sum(base)
    base[-1] += rem
    if widths and widths[0] < widths[1] and n_cols >= 2:
        narrow = widths[0]
        rest = total - narrow
        each = rest // (n_cols - 1)
        base = [narrow] + [each] * (n_cols - 1)
        base[-1] += rest - each * (n_cols - 1)
    return base


def add_styled_table(doc: Document, template: Table, header: list[str], rows: list[list[str]]) -> Table:
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    _clone_tbl_pr(template, table)
    _set_grid(table, _template_col_widths(template, n_cols))

    _set_row_height(table.rows[0], 672)
    for ci, text in enumerate(header):
        cell = table.rows[0].cells[ci]
        _style_cell(cell, HEADER_FILL, header=True)
        _set_cell_text(cell, text, header=True)

    for ri, row in enumerate(rows):
        fill = BODY_FILLS[ri % 2]
        _set_row_height(table.rows[ri + 1], 654)
        for ci in range(n_cols):
            text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri + 1].cells[ci]
            _style_cell(cell, fill, first_body=(ri == 0))
            _set_cell_text(cell, text, header=False)
    return table


def _style_heading(paragraph: Paragraph, level: int) -> None:
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 2")
    try:
        paragraph.style = style_name
    except KeyError:
        paragraph.style = "Normal"
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)


def clear_body_keep_section(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def append_blocks(doc: Document, template: Table, blocks: list[dict]) -> None:
    for block in blocks:
        if block["type"] == "heading":
            p = doc.add_paragraph(block["text"])
            _style_heading(p, block["level"])
        elif block["type"] == "para":
            doc.add_paragraph(block["text"])
        elif block["type"] == "table":
            add_styled_table(doc, template, block["header"], block["rows"])
            doc.add_paragraph("")  # spacing after table


def main() -> None:
    ap = argparse.ArgumentParser(description="Append Markdown into report.docx with styled tables")
    ap.add_argument("markdown", type=Path, help="Source .md file")
    ap.add_argument(
        "-t",
        "--template",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "report.docx",
        help="Template/target docx (default: ../report.docx)",
    )
    ap.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output path (default: overwrite template)",
    )
    ap.add_argument(
        "--replace",
        action="store_true",
        help="Replace document body (style still cloned from first table before clear)",
    )
    args = ap.parse_args()

    md_text = args.markdown.read_text(encoding="utf-8")
    blocks = parse_md(md_text)
    doc = Document(str(args.template))
    if not doc.tables:
        raise SystemExit("template has no table to clone style from")
    # keep a style donor open before optional body clear
    style_donor = Document(str(args.template)).tables[0]

    if args.replace:
        clear_body_keep_section(doc)
    append_blocks(doc, style_donor, blocks)

    out = args.output or args.template
    doc.save(str(out))
    print(f"wrote {out} ({len(blocks)} blocks)")


if __name__ == "__main__":
    main()
